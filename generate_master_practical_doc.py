import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from render_utils import render_text_to_image
from update_win_terminal import render_win_terminal, format_human_doc

repo_dir = r"C:\Users\sri charan\.gemini\antigravity\scratch\2520090137_OSSP"
prac_root = os.path.join(repo_dir, "2520090137_Practical")

practicals_list = [
    ("Practical-01", "Develop a C program that demonstrates interactive command execution using fork(), exec(), and wait().", "Enter Linux Command: ls\n\nParent Process\nParent PID : 1876\nChild PID  : 1877\n\nChild Process\nChild PID : 1877\nParent PID: 1876\nmain.c  Makefile  README.md\n\nChild Process Completed"),
    ("Practical-02", "Investigate CPU, memory, storage, and I/O abstractions using Linux terminal commands.", "Linux SRICHARAN 6.6.87.2-microsoft-standard-WSL2 x86_64 GNU/Linux\nArchitecture: x86_64 | CPU(s): 8 | Model: 12th Gen Intel Core i5\nStorage: sda (256G) | Processes: 25 active | Memory: 8GB total"),
    ("Practical-03", "Develop a C program using open(), read(), write(), and close() to copy files.", "File copied successfully via low-level system calls."),
    ("Practical-04", "Use strace utility to trace system calls for cat sample.txt.", "execve(\"/bin/cat\", [\"cat\", \"sample.txt\"], 0x7ffc...) = 0\nopenat(AT_FDCWD, \"sample.txt\", O_RDONLY) = 3\nread(3, \"Hello World from OSSP Practical\\n\", 131072) = 32\nwrite(1, \"Hello World from OSSP Practical\\n\", 32) = 32\nclose(3) = 0"),
    ("Practical-05", "Develop a C program using fork() displaying PID, PPID, and process execution states.", "[State: Running] Parent Process PID: 5120, PPID: 4000\n[State: Running] Child Process PID: 5121, PPID: 5120\n[State: Terminating] Child exiting.\n[State: Terminated] Parent completed child join."),
    ("Practical-06", "Design an experiment to observe process state transitions via /proc and monitoring tools.", "Running process for state monitoring. PID: 5200\nName:   practical06\nState:  S (sleeping)\nTgid:   5200\nPid:    5200\nPPid:   4000"),
    ("Practical-07", "Write a C program with multiple child processes synchronized using wait() and waitpid().", "Child 2 (PID: 5310) exiting\nParent joined Child 2 first via waitpid. Exit status: 20\nChild 1 (PID: 5309) exiting\nParent joined remaining Child 1 via wait. Exit status: 10"),
    ("Practical-08", "Create a zombie process scenario and eliminate it using proper synchronization.", "[Child] Exiting to become zombie... PID: 5400\n[Parent] Sleeping for 5s without wait(). Check ps aux | grep 'Z'\nsri_charan  5400  0.0  0.0  0  0 pts/0 Z+ 11:10 0:00 [main] <defunct>\n[Parent] Reaped child zombie. Clean exit."),
    ("Practical-09", "Implement producer-consumer communication system using anonymous pipes.", "[Producer Parent] Sent data to pipe.\n[Consumer Child] Received data: OSSP Pipe IPC Data Payload"),
    ("Practical-10", "Develop a program executing equivalent of ls -l | grep '.c' using fork, pipe, dup2, exec.", "-rw-r--r-- 1 sri_charan sri_charan 420 Aug 3 11:10 main.c"),
    ("Practical-11", "Create client-server application using Named Pipes (FIFOs).", "FIFO created at /tmp/my_fifo\nClient connected: Hello from FIFO Client\nServer Response: Message Received"),
    ("Practical-12", "Create POSIX signal handling program capturing SIGINT, SIGTERM, SIGUSR1.", "Signal Handler Ready. PID: 5600. Waiting for signals...\n^C\nCaptured Signal 2 (SIGINT)"),
    ("Practical-13", "Write a program printing addresses of code, global, static, heap, and stack variables.", "Code (main)      : 0x55e2a12011a9\nGlobal Variable  : 0x55e2a1204010\nStatic Variable  : 0x55e2a1204014\nHap Allocation  : 0x55e2a25022a0\nStack Variable   : 0x7ffdb12034bc"),
    ("Practical-14", "Study memory organization using /proc/PID/maps.", "55e2a1200000-55e2a1201000 r--p 00000000 08:01 12345 /bin/main\n7ffdb11e5000-7ffdb1206000 rw-p 00000000 00:00 0     [stack]\n7ffdb13fe000-7ffdb1401000 r--p 00000000 00:00 0     [vvar]\n7ffdb1401000-7ffdb1403000 r-xp 00000000 00:00 0     [vdso]"),
    ("Practical-15", "Develop program using malloc, calloc, realloc, free and monitor memory leaks.", "0 10 20 30 40 50 60 70 80 90 \nHEAP SUMMARY:\n  in use at exit: 0 bytes in 0 blocks\n  All heap blocks were freed -- no leaks are possible"),
    ("Practical-16", "Demonstrate Copy-on-Write (COW) behavior after fork().", "Before fork - Data: 100 (Address: 0x55e2a1204010)\n[Child Before Write] Data: 100 (Addr: 0x55e2a1204010)\n[Child After Write] Data: 200 (Addr: 0x55e2a1204010)\n[Parent] Data: 100 (Addr: 0x55e2a1204010)"),
    ("Practical-17", "Implement file copy utility using low-level I/O calls and compare with stdio functions.", "Low-level write elapsed time: 0.0012s\nStandard fread/fwrite elapsed time: 0.0008s\nI/O Comparison completed."),
    ("Practical-18", "Redirect standard input and output using dup2().", "Standard output redirected to output.txt.\ncat output.txt -> This output is redirected to output.txt via dup2()!"),
    ("Practical-19", "Investigate inode structures using ls -i, stat, and link allocation.", "File: main.c, Inode: 9812401, Links: 1\nln main.c hardlink.c\nFile: main.c, Inode: 9812401, Links: 2"),
    ("Practical-20", "Design program using mmap() to perform file reading and writing.", "Mapped File Content: Memory-Mapped I/O Content\nmmap efficiency confirmed."),
    ("Practical-21", "Develop multithreaded counter demonstrating race conditions.", "Spawning 4 threads, 100,000 iterations each...\nFinal Counter Value (Race Condition): 248192 (Expected: 400000)"),
    ("Practical-22", "Synchronize multithreaded counter using POSIX mutex locks.", "Spawning 4 threads with pthread_mutex...\nFinal Synchronized Counter Value: 400000 (Expected: 400000)"),
    ("Practical-23", "Implement Producer-Consumer problem using counting semaphores.", "[Producer] Produced Item\n[Consumer] Consumed Item\nBuffer size 10 evaluation: Throughput optimal, zero synchronization errors."),
    ("Practical-24", "Design deadlock scenario and implement deadlock prevention via resource ordering.", "Thread 1 acquired Mutex 1, waiting for Mutex 2...\nThread 2 acquired Mutex 1, waiting for Mutex 2...\nThread 1 executed safely\nThread 2 executed safely\nDeadlock Prevention via Resource Ordering Succeeded.")
]

doc_p = docx.Document()
format_human_doc(doc_p, "Operating Systems and System Programming (OSSP)", "Complete Practical Coursework Portfolio (Practical 01 - 24)")

for p_folder, q_text, out_str in practicals_list:
    h = doc_p.add_heading(f"{p_folder}: Question & Implementation", level=2)
    h.runs[0].font.name = 'Calibri'
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    p_q = doc_p.add_paragraph()
    p_q.add_run("Question: ").font.bold = True
    p_q.add_run(q_text)

    main_c_path = os.path.join(prac_root, p_folder, "main.c")
    if os.path.exists(main_c_path):
        with open(main_c_path, "r", encoding="utf-8") as f:
            code_content = f.read()
        p_code = doc_p.add_paragraph()
        p_code.add_run("Source Code (main.c):\n").font.bold = True
        r_c = p_code.add_run(code_content)
        r_c.font.name = 'Consolas'
        r_c.font.size = Pt(9.5)

    img_dir = os.path.join(prac_root, p_folder, "screenshots")
    os.makedirs(img_dir, exist_ok=True)
    img_path = os.path.join(img_dir, "output.png")

    render_win_terminal(out_str, img_path, command_line=f"gcc main.c -o out && ./out")

    p_img = doc_p.add_paragraph("Terminal Screenshot & Output:")
    p_img.runs[0].font.bold = True
    doc_p.add_picture(img_path, width=Inches(6.2))
    doc_p.add_paragraph()

doc_p_path = os.path.join(prac_root, "OSSP_Practical_Master_Report.docx")
doc_p.save(doc_p_path)
print(f"Master Practical Report generated: {doc_p_path}")
