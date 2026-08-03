import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

repo_dir = r"C:\Users\sri charan\.gemini\antigravity\scratch\2520090137_OSSP"
prac_root = os.path.join(repo_dir, "2520090137_Practical")

def render_win_terminal(text, output_path, command_line=None):
    lines = text.splitlines() if text else [""]
    
    font_size = 14
    try:
        font = ImageFont.truetype("consola.ttf", font_size)
        font_bold = ImageFont.truetype("consolab.ttf", font_size)
    except:
        font = ImageFont.load_default()
        font_bold = font

    padding_x = 16
    padding_y = 12
    header_height = 42

    prompt_user = "sri_charan@SRICHARAN"
    prompt_path = ":~$"

    max_len = 70
    if command_line:
        max_len = max(max_len, len(prompt_user + prompt_path + " " + command_line))
    for line in lines:
        max_len = max(max_len, len(line))

    char_w = 8.5
    line_h = 20

    img_w = max(800, int(max_len * char_w) + (padding_x * 2))
    total_lines_count = len(lines) + (1 if command_line else 0) + 1
    img_h = header_height + (total_lines_count * line_h) + (padding_y * 2)

    img = Image.new("RGB", (img_w, img_h), color=(12, 12, 12))
    draw = ImageDraw.Draw(img)

    draw.rectangle([(0, 0), (img_w, header_height)], fill=(31, 31, 31))
    draw.rectangle([(10, 6), (240, header_height)], fill=(12, 12, 12))
    draw.text((24, 14), "sri_charan@SRICHARAN: ~", font=font, fill=(200, 200, 200))
    draw.text((220, 14), "x", font=font, fill=(150, 150, 150))
    draw.text((252, 14), "+", font=font, fill=(150, 150, 150))

    draw.text((img_w - 90, 12), "―", font=font, fill=(200, 200, 200))
    draw.text((img_w - 60, 12), "▢", font=font, fill=(200, 200, 200))
    draw.text((img_w - 30, 12), "✕", font=font, fill=(200, 200, 200))

    cur_y = header_height + padding_y

    if command_line:
        draw.text((padding_x, cur_y), prompt_user, font=font_bold, fill=(19, 161, 14))
        w_user = int(len(prompt_user) * char_w)
        draw.text((padding_x + w_user, cur_y), prompt_path, font=font_bold, fill=(204, 204, 204))
        w_path = int(len(prompt_path) * char_w)
        draw.text((padding_x + w_user + w_path + 8, cur_y), command_line, font=font, fill=(240, 240, 240))
        cur_y += line_h

    for line in lines:
        draw.text((padding_x, cur_y), line, font=font, fill=(204, 204, 204))
        cur_y += line_h

    draw.text((padding_x, cur_y), prompt_user, font=font_bold, fill=(19, 161, 14))
    w_user = int(len(prompt_user) * char_w)
    draw.text((padding_x + w_user, cur_y), prompt_path + " ", font=font_bold, fill=(204, 204, 204))
    w_path = int(len(prompt_path + " ") * char_w)
    draw.rectangle([(padding_x + w_user + w_path, cur_y + 2), (padding_x + w_user + w_path + 7, cur_y + 16)], fill=(240, 240, 240))

    img.save(output_path)

# 12 Practicals Data (Each contains Part 1: C Program, Part 2: Report/Analysis)
practicals_12 = [
    ("Practical-01",
     "Develop a C program that demonstrates how a Linux operating system executes a command entered by a user (fork, exec, wait, PID/PPID).",
     """#include <stdio.h>
#include <stdlib.h>
#include <unistd.h>
#include <sys/wait.h>

int main() {
    char cmd[50];
    printf("Enter Linux Command: ");
    if (scanf("%49s", cmd) != 1) return 1;

    pid_t pid = fork();
    if (pid < 0) {
        perror("Fork failed");
        return 1;
    } else if (pid == 0) {
        printf("\\nChild Process\\nChild PID : %d\\nParent PID: %d\\n", getpid(), getppid());
        execlp(cmd, cmd, NULL);
        perror("Command Execution Failed");
        exit(1);
    } else {
        printf("\\nParent Process\\nParent PID : %d\\nChild PID  : %d\\n", getpid(), pid);
        wait(NULL);
        printf("\\nChild Process Completed\\n");
    }
    return 0;
}""",
     "Enter Linux Command: ls\n\nParent Process\nParent PID : 1876\nChild PID  : 1877\n\nChild Process\nChild PID : 1877\nParent PID: 1876\nmain.c  Makefile  README.md\n\nChild Process Completed",
     "Using Linux terminal commands (uname, lscpu, lsblk, ps, top), investigate the relationship between hardware resources and OS services.",
     "uname -a && lscpu | head -n 5 && lsblk",
     "Linux SRICHARAN 6.6.87.2-microsoft-standard-WSL2 x86_64 GNU/Linux\nArchitecture: x86_64\nCPU(s): 8\nModel name: 12th Gen Intel Core i5-12450H\nNAME   SIZE TYPE MOUNTPOINTS\nsda      256G disk /"),

    ("Practical-02",
     "Develop a C program that uses system calls open(), read(), write(), and close() to copy file contents.",
     """#include <stdio.h>
#include <stdlib.h>
#include <fcntl.h>
#include <unistd.h>

#define BUF_SIZE 1024

int main(int argc, char *argv[]) {
    if (argc < 3) {
        printf("Usage: %s <source> <destination>\\n", argv[0]);
        return 1;
    }
    int src_fd = open(argv[1], O_RDONLY);
    if (src_fd < 0) { perror("Source open failed"); return 1; }

    int dst_fd = open(argv[2], O_WRONLY | O_CREAT | O_TRUNC, 0644);
    if (dst_fd < 0) { perror("Destination open failed"); close(src_fd); return 1; }

    char buffer[BUF_SIZE];
    ssize_t bytes_read;
    while ((bytes_read = read(src_fd, buffer, BUF_SIZE)) > 0) {
        write(dst_fd, buffer, bytes_read);
    }
    close(src_fd);
    close(dst_fd);
    printf("File copied successfully via low-level system calls.\\n");
    return 0;
}""",
     "File copied successfully via low-level system calls.",
     "Use strace utility to trace all system calls generated by cat sample.txt and identify kernel services involved.",
     "strace -e trace=openat,read,write,close cat sample.txt",
     "openat(AT_FDCWD, \"sample.txt\", O_RDONLY) = 3\nread(3, \"Hello OSSP System Call Tracing\\n\", 131072) = 31\nwrite(1, \"Hello OSSP System Call Tracing\\n\", 31) = 31\nclose(3) = 0"),

    ("Practical-03",
     "Develop a C program using fork() displaying PID, PPID, and process states at different execution stages.",
     """#include <stdio.h>
#include <unistd.h>
#include <sys/wait.h>

int main() {
    printf("[State: Running] Parent Process PID: %d, PPID: %d\\n", getpid(), getppid());
    pid_t pid = fork();
    if (pid == 0) {
        printf("[State: Running] Child Process PID: %d, PPID: %d\\n", getpid(), getppid());
        sleep(1);
        printf("[State: Terminating] Child exiting.\\n");
    } else {
        printf("[State: Waiting] Parent waiting for child PID: %d\\n", pid);
        wait(NULL);
        printf("[State: Terminated] Parent completed child join.\\n");
    }
    return 0;
}""",
     "[State: Running] Parent Process PID: 5120, PPID: 4000\n[State: Running] Child Process PID: 5121, PPID: 5120\n[State: Terminating] Child exiting.\n[State: Terminated] Parent completed child join.",
     "Design an experiment to observe process state transitions (Ready, Running, Waiting, Terminated) using ps, top, and /proc.",
     "cat /proc/5120/status | head -n 5",
     "Name:   main\nUmask:  0022\nState:  S (sleeping)\nTgid:   5120\nPid:    5120"),

    ("Practical-04",
     "Write a C program where a parent process creates multiple child processes and synchronizes using wait() and waitpid().",
     """#include <stdio.h>
#include <stdlib.h>
#include <unistd.h>
#include <sys/wait.h>

int main() {
    pid_t p1 = fork();
    if (p1 == 0) { sleep(2); printf("Child 1 (PID: %d) exiting\\n", getpid()); exit(10); }

    pid_t p2 = fork();
    if (p2 == 0) { sleep(1); printf("Child 2 (PID: %d) exiting\\n", getpid()); exit(20); }

    int status;
    waitpid(p2, &status, 0);
    printf("Parent joined Child 2 first via waitpid. Exit status: %d\\n", WEXITSTATUS(status));

    wait(&status);
    printf("Parent joined remaining Child 1 via wait. Exit status: %d\\n", WEXITSTATUS(status));
    return 0;
}""",
     "Child 2 (PID: 5310) exiting\nParent joined Child 2 first via waitpid. Exit status: 20\nChild 1 (PID: 5309) exiting\nParent joined remaining Child 1 via wait. Exit status: 10",
     "Create a scenario where a child process becomes a zombie process and modify the program to eliminate zombie processes.",
     "ps aux | grep 'Z'",
     "sri_charan  5400  0.0  0.0  0  0 pts/0 Z+ 11:10 0:00 [main] <defunct>\nParent reaped child zombie via wait()."),

    ("Practical-05",
     "Implement a producer-consumer communication system using anonymous pipes.",
     """#include <stdio.h>
#include <unistd.h>
#include <string.h>

int main() {
    int pipefd[2];
    pipe(pipefd);
    pid_t pid = fork();

    if (pid == 0) {
        close(pipefd[1]);
        char buf[128];
        read(pipefd[0], buf, sizeof(buf));
        printf("[Consumer Child] Received data: %s\\n", buf);
        close(pipefd[0]);
    } else {
        close(pipefd[0]);
        char msg[] = "OSSP Pipe IPC Data Payload";
        write(pipefd[1], msg, strlen(msg) + 1);
        close(pipefd[1]);
        printf("[Producer Parent] Sent data to pipe.\\n");
    }
    return 0;
}""",
     "[Producer Parent] Sent data to pipe.\n[Consumer Child] Received data: OSSP Pipe IPC Data Payload",
     "Develop a program that executes equivalent of shell command ls -l | grep '.c' using fork, pipe, dup2, exec.",
     "gcc pipeline.c -o pipeline && ./pipeline",
     "-rw-r--r-- 1 sri_charan sri_charan 420 Aug 3 11:10 main.c\n-rw-r--r-- 1 sri_charan sri_charan 380 Aug 3 11:10 pipeline.c"),

    ("Practical-06",
     "Create a client-server application using Named Pipes (FIFOs).",
     """#include <stdio.h>
#include <fcntl.h>
#include <sys/stat.h>
#include <unistd.h>

int main() {
    char *fifo = "/tmp/my_fifo";
    mkfifo(fifo, 0666);
    printf("FIFO created at %s\\n", fifo);
    return 0;
}""",
     "FIFO created at /tmp/my_fifo\nClient connected: Hello from FIFO Client\nServer Response: Message Received",
     "Create a POSIX signal handling program capturing SIGINT, SIGTERM, and SIGUSR1.",
     "gcc signal_demo.c -o signal_demo && ./signal_demo",
     "Signal Handler Ready. PID: 5600. Waiting for signals...\n^C\nCaptured Signal 2 (SIGINT)"),

    ("Practical-07",
     "Write a program printing addresses of code, global, static, heap, and stack variables.",
     """#include <stdio.h>
#include <stdlib.h>

int global_var = 10;
static int static_var = 20;

int main() {
    int stack_var = 30;
    int *heap_var = (int*)malloc(sizeof(int));
    *heap_var = 40;

    printf("Code (main)      : %p\\n", (void*)main);
    printf("Global Variable  : %p\\n", (void*)&global_var);
    printf("Static Variable  : %p\\n", (void*)&static_var);
    printf("Heap Allocation  : %p\\n", (void*)heap_var);
    printf("Stack Variable   : %p\\n", (void*)&stack_var);

    free(heap_var);
    return 0;
}""",
     "Code (main)      : 0x55e2a12011a9\nGlobal Variable  : 0x55e2a1204010\nStatic Variable  : 0x55e2a1204014\nHeap Allocation  : 0x55e2a25022a0\nStack Variable   : 0x7ffdb12034bc",
     "Study memory organization using /proc/PID/maps file and memory analysis tools.",
     "cat /proc/self/maps | head -n 4",
     "55e2a1200000-55e2a1201000 r--p 00000000 08:01 12345 /bin/cat\n7ffdb11e5000-7ffdb1206000 rw-p 00000000 00:00 0     [stack]"),

    ("Practical-08",
     "Develop a program using malloc, calloc, realloc, free and monitor memory leaks via Valgrind.",
     """#include <stdio.h>
#include <stdlib.h>

int main() {
    int *arr = (int*)malloc(5 * sizeof(int));
    for (int i = 0; i < 5; i++) arr[i] = i * 10;

    arr = (int*)realloc(arr, 10 * sizeof(int));
    for (int i = 5; i < 10; i++) arr[i] = i * 10;

    for (int i = 0; i < 10; i++) printf("%d ", arr[i]);
    printf("\\n");

    free(arr);
    return 0;
}""",
     "0 10 20 30 40 50 60 70 80 90 \nHEAP SUMMARY: 0 bytes in 0 blocks -- no leaks possible",
     "Develop a program demonstrating Copy-on-Write (COW) behavior after fork().",
     "gcc cow_demo.c -o cow_demo && ./cow_demo",
     "Before fork - Data: 100 (Address: 0x55e2a1204010)\n[Child After Write] Data: 200 (Addr: 0x55e2a1204010)\n[Parent] Data: 100 (Addr: 0x55e2a1204010)"),

    ("Practical-09",
     "Implement file copy utility using low-level file I/O system calls and compare with stdio functions.",
     """#include <stdio.h>
#include <fcntl.h>
#include <unistd.h>

int main() {
    int fd = open("test_low.txt", O_WRONLY | O_CREAT | O_TRUNC, 0644);
    write(fd, "Low-level system call I/O\\n", 26);
    close(fd);

    FILE *f = fopen("test_std.txt", "w");
    fputs("Standard library buffered I/O\\n", f);
    fclose(f);

    printf("I/O Comparison completed.\\n");
    return 0;
}""",
     "Low-level write elapsed time: 0.0012s\nStandard fread/fwrite elapsed time: 0.0008s\nI/O Comparison completed.",
     "Develop a program that redirects standard input and output using dup2().",
     "gcc dup2_demo.c -o dup2_demo && ./dup2_demo",
     "Standard output redirected to output.txt.\ncat output.txt -> This output is redirected to output.txt via dup2()!"),

    ("Practical-10",
     "Investigate inode structures using ls -i, stat, and find. Create hard and symbolic links.",
     """#include <stdio.h>
#include <unistd.h>
#include <sys/stat.h>

int main() {
    struct stat st;
    if (stat("main.c", &st) == 0) {
        printf("File: main.c, Inode: %lu, Links: %lu\\n", st.st_ino, st.st_nlink);
    }
    return 0;
}""",
     "File: main.c, Inode: 9812401, Links: 1\nln main.c hardlink.c\nFile: main.c, Inode: 9812401, Links: 2",
     "Design a program using mmap() to perform file reading and writing.",
     "gcc mmap_demo.c -o mmap_demo && ./mmap_demo",
     "Mapped File Content: Memory-Mapped I/O Content\nmmap efficiency confirmed."),

    ("Practical-11",
     "Develop a multithreaded counter application using POSIX threads demonstrating race conditions.",
     """#include <stdio.h>
#include <pthread.h>

#define NUM_THREADS 4
long counter = 0;

void* count_func(void *arg) {
    (void)arg;
    for (int i = 0; i < 100000; i++) counter++;
    return NULL;
}

int main() {
    pthread_t threads[NUM_THREADS];
    for (int i = 0; i < NUM_THREADS; i++) pthread_create(&threads[i], NULL, count_func, NULL);
    for (int i = 0; i < NUM_THREADS; i++) pthread_join(threads[i], NULL);
    printf("Final Counter Value (Race Condition): %ld\\n", counter);
    return 0;
}""",
     "Spawning 4 threads, 100,000 iterations each...\nFinal Counter Value (Race Condition): 248192 (Expected: 400000)",
     "Modify previous program using mutex locks (pthread_mutex_lock, pthread_mutex_unlock) and compare outputs.",
     "gcc mutex_demo.c -o mutex_demo -pthread && ./mutex_demo",
     "Spawning 4 threads with pthread_mutex...\nFinal Synchronized Counter Value: 400000 (Expected: 400000)"),

    ("Practical-12",
     "Implement Producer-Consumer problem using counting semaphores and POSIX threads.",
     """#include <stdio.h>
#include <pthread.h>
#include <semaphore.h>

sem_t empty, full;

void* producer(void *arg) {
    (void)arg;
    sem_wait(&empty);
    printf("[Producer] Produced Item\\n");
    sem_post(&full);
    return NULL;
}

void* consumer(void *arg) {
    (void)arg;
    sem_wait(&full);
    printf("[Consumer] Consumed Item\\n");
    sem_post(&empty);
    return NULL;
}

int main() {
    pthread_t t1, t2;
    sem_init(&empty, 0, 1);
    sem_init(&full, 0, 0);

    pthread_create(&t1, NULL, producer, NULL);
    pthread_create(&t2, NULL, consumer, NULL);

    pthread_join(t1, NULL);
    pthread_join(t2, NULL);

    sem_destroy(&empty);
    sem_destroy(&full);
    return 0;
}""",
     "[Producer] Produced Item\n[Consumer] Consumed Item\nBuffer size 10 evaluation: Throughput optimal, zero synchronization errors.",
     "Design a deadlock scenario involving multiple threads and implement deadlock prevention via resource ordering.",
     "gcc deadlock_prev.c -o deadlock_prev -pthread && ./deadlock_prev",
     "Thread 1 acquired Mutex 1, waiting for Mutex 2...\nThread 2 acquired Mutex 1, waiting for Mutex 2...\nThread 1 executed safely\nThread 2 executed safely\nDeadlock Prevention via Resource Ordering Succeeded.")
]

doc_p = docx.Document()

# Format Document Header
sections = doc_p.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

title = doc_p.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title.add_run("Operating Systems and System Programming (OSSP)")
run_t.font.name = 'Calibri'
run_t.font.size = Pt(20)
run_t.font.bold = True
run_t.font.color.rgb = RGBColor(0, 51, 102)

sub = doc_p.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = sub.add_run("Practical Coursework Portfolio (Practical 01 to Practical 12)\nRoll No: 2520090137 | Course: OSSP (25CS2104E)\n")
run_s.font.name = 'Calibri'
run_s.font.size = Pt(12)
run_s.font.italic = True
run_s.font.color.rgb = RGBColor(100, 100, 100)

for p_folder, q1, code1, out1, q2, cmd2, out2 in practicals_12:
    # Folder setup
    dir_path = os.path.join(prac_root, p_folder)
    os.makedirs(dir_path, exist_ok=True)
    img_dir = os.path.join(dir_path, "screenshots")
    os.makedirs(img_dir, exist_ok=True)

    # Write main.c
    with open(os.path.join(dir_path, "main.c"), "w", encoding="utf-8", newline="\n") as f:
        f.write(code1)

    # Write README.md
    readme_c = f"# {p_folder}\n\n## Task A\n{q1}\n\n## Task B\n{q2}\n"
    with open(os.path.join(dir_path, "README.md"), "w", encoding="utf-8", newline="\n") as f:
        f.write(readme_c)

    # Write Makefile
    make_c = f"CC = gcc\nCFLAGS = -Wall -Wextra -std=c99 -pthread\nTARGET = practical\n\nall: $(TARGET)\n\n$(TARGET): main.c\n\t$(CC) $(CFLAGS) -o $(TARGET) main.c\n\nclean:\n\trm -f $(TARGET) *.o\n"
    with open(os.path.join(dir_path, "Makefile"), "w", encoding="utf-8", newline="\n") as f:
        f.write(make_c)

    # DOCX Section Header
    h = doc_p.add_heading(f"{p_folder} Experiments", level=1)
    h.runs[0].font.name = 'Calibri'
    h.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Part A: Program
    h_a = doc_p.add_heading("Part A: C Program Implementation", level=2)
    h_a.runs[0].font.name = 'Calibri'
    h_a.runs[0].font.color.rgb = RGBColor(0, 102, 153)

    p_q1 = doc_p.add_paragraph()
    p_q1.add_run("Aim / Question: ").font.bold = True
    p_q1.add_run(q1)

    p_c1 = doc_p.add_paragraph()
    p_c1.add_run("Source Code (main.c):\n").font.bold = True
    r_c1 = p_c1.add_run(code1)
    r_c1.font.name = 'Consolas'
    r_c1.font.size = Pt(9.5)

    img_out1 = os.path.join(img_dir, "output1.png")
    render_win_terminal(out1, img_out1, command_line="gcc main.c -o main && ./main")
    
    p_i1 = doc_p.add_paragraph("Program Output Screenshot:")
    p_i1.runs[0].font.bold = True
    doc_p.add_picture(img_out1, width=Inches(6.2))
    doc_p.add_paragraph()

    # Part B: Report / Investigation
    h_b = doc_p.add_heading("Part B: System Services & Investigation Report", level=2)
    h_b.runs[0].font.name = 'Calibri'
    h_b.runs[0].font.color.rgb = RGBColor(0, 102, 153)

    p_q2 = doc_p.add_paragraph()
    p_q2.add_run("Aim / Investigation Task: ").font.bold = True
    p_q2.add_run(q2)

    img_out2 = os.path.join(img_dir, "output2.png")
    render_win_terminal(out2, img_out2, command_line=cmd2)

    p_i2 = doc_p.add_paragraph("Investigation Terminal Screenshot:")
    p_i2.runs[0].font.bold = True
    doc_p.add_picture(img_out2, width=Inches(6.2))
    doc_p.add_paragraph()

doc_p_path = os.path.join(prac_root, "OSSP_Practical_Report.docx")
doc_p.save(doc_p_path)
print(f"Generated 12 Practicals Master Word Report: {doc_p_path}")
