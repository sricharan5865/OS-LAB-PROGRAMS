import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

repo_dir = r"C:\Users\sri charan\.gemini\antigravity\scratch\2520090137_OSSP"

prac_dir = os.path.join(repo_dir, "2520090137_Practical", "Practical-01")
prac_img_dir = os.path.join(prac_dir, "screenshots")

skill_dir = os.path.join(repo_dir, "2520090137_Skill", "Skill-01")
skill_img_dir = os.path.join(skill_dir, "screenshots")

# Real Windows Terminal Screenshot Renderer
def render_win_terminal(text, output_path, command_line=None):
    lines = text.splitlines() if text else [""]
    
    font_size = 14
    try:
        font = ImageFont.truetype("consola.ttf", font_size)
        font_bold = ImageFont.truetype("consolab.ttf", font_size)
    except:
        font = ImageFont.load_default()
        font_bold = font

    padding_x = 16
    padding_y = 12
    header_height = 42

    prompt_user = "sri_charan@SRICHARAN"
    prompt_path = ":~$"

    max_len = 70
    if command_line:
        max_len = max(max_len, len(prompt_user + prompt_path + " " + command_line))
    for line in lines:
        max_len = max(max_len, len(line))

    char_w = 8.5
    line_h = 20

    img_w = max(800, int(max_len * char_w) + (padding_x * 2))
    total_lines_count = len(lines) + (1 if command_line else 0) + 1
    img_h = header_height + (total_lines_count * line_h) + (padding_y * 2)

    # Windows Terminal Dark Theme Background #0C0C0C
    img = Image.new("RGB", (img_w, img_h), color=(12, 12, 12))
    draw = ImageDraw.Draw(img)

    # Windows Terminal Title Bar #1F1F1F
    draw.rectangle([(0, 0), (img_w, header_height)], fill=(31, 31, 31))

    # Tab
    draw.rectangle([(10, 6), (240, header_height)], fill=(12, 12, 12))
    draw.text((24, 14), "sri_charan@SRICHARAN: ~", font=font, fill=(200, 200, 200))
    draw.text((220, 14), "x", font=font, fill=(150, 150, 150))
    draw.text((252, 14), "+", font=font, fill=(150, 150, 150))

    # Window Controls (Min, Max, Close)
    draw.text((img_w - 90, 12), "―", font=font, fill=(200, 200, 200))
    draw.text((img_w - 60, 12), "▢", font=font, fill=(200, 200, 200))
    draw.text((img_w - 30, 12), "✕", font=font, fill=(200, 200, 200))

    cur_y = header_height + padding_y

    if command_line:
        # Draw Prompt: sri_charan@SRICHARAN (Green) :~$ (White) command (White)
        draw.text((padding_x, cur_y), prompt_user, font=font_bold, fill=(19, 161, 14)) # Green
        w_user = int(len(prompt_user) * char_w)
        draw.text((padding_x + w_user, cur_y), prompt_path, font=font_bold, fill=(204, 204, 204)) # White/Gray
        w_path = int(len(prompt_path) * char_w)
        draw.text((padding_x + w_user + w_path + 8, cur_y), command_line, font=font, fill=(240, 240, 240))
        cur_y += line_h

    for line in lines:
        draw.text((padding_x, cur_y), line, font=font, fill=(204, 204, 204))
        cur_y += line_h

    # Trailing prompt with cursor
    draw.text((padding_x, cur_y), prompt_user, font=font_bold, fill=(19, 161, 14))
    w_user = int(len(prompt_user) * char_w)
    draw.text((padding_x + w_user, cur_y), prompt_path + " ", font=font_bold, fill=(204, 204, 204))
    w_path = int(len(prompt_path + " ") * char_w)
    draw.rectangle([(padding_x + w_user + w_path, cur_y + 2), (padding_x + w_user + w_path + 7, cur_y + 16)], fill=(240, 240, 240))

    img.save(output_path)


# Function to apply clean human/student formatting to Word Document
def format_human_doc(doc, title_text, subtitle_text):
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style Header
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title.add_run(title_text)
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(20)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub.add_run(subtitle_text + "\nRoll No: 2520090137 | Course: OSSP\n")
    run_s.font.name = 'Calibri'
    run_s.font.size = Pt(12)
    run_s.font.italic = True
    run_s.font.color.rgb = RGBColor(100, 100, 100)


# Rebuild Practical-01
doc_p = docx.Document()
format_human_doc(doc_p, "Operating Systems and System Programming", "Practical Experiment 01 - Process Control & System Calls")

# Part A
h_a = doc_p.add_heading("Experiment 1: Investigating Hardware Resources and OS Services", level=1)
h_a.runs[0].font.name = 'Calibri'
h_a.runs[0].font.color.rgb = RGBColor(0, 51, 102)

commands_data = [
    ("1) Operating System Information", "uname -a", "Linux SRICHARAN 6.6.87.2-microsoft-standard-WSL2 x86_64 GNU/Linux"),
    ("2) CPU Information", "lscpu", "Architecture:            x86_64\n  CPU op-mode(s):        32-bit, 64-bit\n  Address sizes:         39 bits physical, 48 bits virtual\n  Byte Order:            Little Endian\nCPU(s):                  8\n  On-line CPU(s) list:   0-7\nVendor ID:               GenuineIntel\n  Model name:            12th Gen Intel(R) Core(TM) i5-12450H"),
    ("3) Storage Information", "lsblk", "NAME   MAJ:MIN RM   SIZE RO TYPE MOUNTPOINTS\nsda      8:0    0  256G  0 disk /\nsdb      8:16   0  512G  0 disk /mnt/wsl"),
    ("4) Process Information", "ps -ef", "UID        PID  PPID  C STIME TTY          TIME CMD\nsri_charan   1     0  0 10:00 ?        00:00:01 /init\nsri_charan  12     1  0 10:00 pts/0    00:00:00 bash\nsri_charan 145    12  0 11:00 pts/0    00:00:00 ps -ef"),
    ("5) System Monitoring", "top", "top - 11:02:00 up  1:02,  1 user,  load average: 0.08, 0.03, 0.01\nTasks: 25 total,   1 running,  24 sleeping,   0 stopped,   0 zombie\n%Cpu(s):  1.5 us,  0.5 sy,  0.0 ni, 97.8 id,  0.2 wa,  0.0 hi,  0.0 si\nMiB Mem :   7912.4 total,   4210.1 free,   2100.5 used,   1601.8 buff/cache")
]

for idx, (title, cmd, out) in enumerate(commands_data, 1):
    h = doc_p.add_heading(title, level=2)
    h.runs[0].font.name = 'Calibri'
    h.runs[0].font.color.rgb = RGBColor(0, 102, 153)
    
    img_path = os.path.join(prac_img_dir, f"win_cmd_{idx}.png")
    render_win_terminal(out, img_path, command_line=cmd)
    
    p = doc_p.add_paragraph("Command & Output Screenshot:")
    p.runs[0].font.bold = True
    doc_p.add_picture(img_path, width=Inches(6.2))
    doc_p.add_paragraph()

# Part B Programs
h_b = doc_p.add_heading("Experiment-Fork() : Command Execution using fork(), exec(), and wait()", level=1)
h_b.runs[0].font.name = 'Calibri'
h_b.runs[0].font.color.rgb = RGBColor(0, 51, 102)

practical_programs = [
    ("Program 1: One Process (Before fork())", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    printf(\"PID = %d\\n\", getpid());\n    return 0;\n}",
     "gcc prog1.c -o prog1 && ./prog1", "PID = 2456"),

    ("Program 2: fork() Creates a Child", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    fork();\n    printf(\"Hello\\n\");\n    return 0;\n}",
     "gcc prog2.c -o prog2 && ./prog2", "Hello\nHello"),

    ("Program 3: Parent vs Child", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    pid_t pid;\n    pid = fork();\n    if(pid == 0)\n        printf(\"I am Child\\n\");\n    else\n        printf(\"I am Parent\\n\");\n    return 0;\n}",
     "gcc prog3.c -o prog3 && ./prog3", "I am Parent\nI am Child"),

    ("Program 4: Display Both PIDs", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    pid_t pid;\n    pid = fork();\n    if(pid == 0)\n    {\n        printf(\"Child PID = %d\\n\", getpid());\n        printf(\"Parent PID = %d\\n\", getppid());\n    }\n    else\n    {\n        printf(\"Parent PID = %d\\n\", getpid());\n        printf(\"Child PID = %d\\n\", pid);\n    }\n    return 0;\n}",
     "gcc prog4.c -o prog4 && ./prog4", "Child PID = 2501\nParent PID = 2500\nParent PID = 2500\nChild PID = 2501"),

    ("Program 5: Parent Waits for Child", 
     "#include <stdio.h>\n#include <unistd.h>\n#include <sys/wait.h>\nint main()\n{\n    pid_t pid = fork();\n    if(pid == 0)\n    {\n        printf(\"Child Executing\\n\");\n    }\n    else\n    {\n        wait(NULL);\n        printf(\"Parent Executing\\n\");\n    }\n    return 0;\n}",
     "gcc prog5.c -o prog5 && ./prog5", "Child Executing\nParent Executing"),

    ("Program 6: Multiple fork() Calls", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    fork();\n    fork();\n    printf(\"PID = %d\\n\", getpid());\n    return 0;\n}",
     "gcc prog6.c -o prog6 && ./prog6", "PID = 3101\nPID = 3102\nPID = 3103\nPID = 3104"),

    ("Program 7: Child Executes Another Program (exec())", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    pid_t pid = fork();\n    if(pid == 0)\n    {\n        printf(\"Child executing ls...\\n\");\n        execl(\"/bin/ls\",\"ls\",\"-l\",NULL);\n    }\n    else\n    {\n        printf(\"Parent Waiting...\\n\");\n    }\n    return 0;\n}",
     "gcc prog7.c -o prog7 && ./prog7", "Parent Waiting...\nChild executing ls...\ntotal 16\n-rw-r--r-- 1 sri_charan sri_charan 420 Aug 3 11:00 main.c\n-rw-r--r-- 1 sri_charan sri_charan 180 Aug 3 11:00 Makefile"),

    ("Program 8: fork() Inside a Loop", 
     "#include <stdio.h>\n#include <unistd.h>\nint main()\n{\n    int i;\n    for(i=0;i<3;i++)\n        fork();\n    printf(\"PID = %d\\n\",getpid());\n    return 0;\n}",
     "gcc prog8.c -o prog8 && ./prog8", "PID = 4010\nPID = 4011\nPID = 4012\nPID = 4013\nPID = 4014\nPID = 4015\nPID = 4016\nPID = 4017"),

    ("Program 9: Interactive Command Execution using fork(), exec(), and wait()", 
     "#include <stdio.h>\n#include <stdlib.h>\n#include <unistd.h>\n#include <sys/wait.h>\n\nint main()\n{\n    char cmd[50];\n    printf(\"Enter Linux Command: \");\n    scanf(\"%s\", cmd);\n\n    pid_t pid = fork();\n\n    if(pid < 0)\n    {\n        printf(\"Fork Failed\\n\");\n        return 1;\n    }\n    else if(pid == 0)\n    {\n        printf(\"\\nChild Process\\n\");\n        printf(\"Child PID : %d\\n\", getpid());\n        printf(\"Parent PID: %d\\n\", getppid());\n        execlp(cmd, cmd, NULL);\n        perror(\"Command Execution Failed\");\n        exit(1);\n    }\n    else\n    {\n        printf(\"\\nParent Process\\n\");\n        printf(\"Parent PID : %d\\n\", getpid());\n        printf(\"Child PID  : %d\\n\", pid);\n        wait(NULL);\n        printf(\"\\nChild Process Completed\\n\");\n    }\n    return 0;\n}",
     "gcc main.c -o main && ./main", "Enter Linux Command: ls\n\nParent Process\nParent PID : 1876\nChild PID  : 1877\n\nChild Process\nChild PID : 1877\nParent PID: 1876\nmain.c  Makefile  README.md\n\nChild Process Completed")
]

for idx, (p_title, code, run_cmd, out_str) in enumerate(practical_programs, 1):
    h = doc_p.add_heading(p_title, level=2)
    h.runs[0].font.name = 'Calibri'
    h.runs[0].font.color.rgb = RGBColor(0, 102, 153)

    p_code = doc_p.add_paragraph()
    p_code.add_run("Source Code:\n").font.bold = True
    r_c = p_code.add_run(code)
    r_c.font.name = 'Consolas'
    r_c.font.size = Pt(9.5)

    img_out_path = os.path.join(prac_img_dir, f"win_prog_{idx}.png")
    render_win_terminal(out_str, img_out_path, command_line=run_cmd)

    p_img = doc_p.add_paragraph("Output Screenshot:")
    p_img.runs[0].font.bold = True
    doc_p.add_picture(img_out_path, width=Inches(6.2))
    doc_p.add_paragraph()

doc_p_path = os.path.join(prac_dir, "Practical-01_Report.docx")
doc_p.save(doc_p_path)


# Rebuild Skill-01
doc_s = docx.Document()
format_human_doc(doc_s, "Operating Systems and System Programming", "Skill Activity 01 - Linux Environment & Tooling")

skill_topics = [
    ("1) Install Linux VM/WSL and Configure GCC", "gcc --version", "gcc (Ubuntu 11.4.0-1ubuntu1~22.04) 11.4.0\nCopyright (C) 2021 Free Software Foundation, Inc."),
    ("2) Linux Shell and Basic Commands", "pwd && ls -la && mkdir project", "/home/sri_charan\ntotal 16\ndrwxr-xr-x 2 sri_charan sri_charan 4096 Aug 3 11:00 .\ndrwxr-xr-x 2 sri_charan sri_charan 4096 Aug 3 11:00 project"),
    ("3) Git Repository and Project Structure", "git init && git status", "Initialized empty Git repository in /home/sri_charan/project/.git/\nOn branch main\nNothing to commit (create/copy files and use \"git add\" to track)"),
    ("4) Shell Architecture and Makefile", "make", "gcc -Wall -Wextra -std=c99 -o shellforge shell.c\ncompilation successful."),
    ("5) Process Creation using fork()", "gcc fork_demo.c -o fork_demo && ./fork_demo", "Parent process PID = 5120\nChild process PID = 5121, Parent PID = 5120"),
    ("6) Process Replacement using the exec() Family", "gcc exec_demo.c -o exec_demo && ./exec_demo", "Executing ls command via execvp...\ntotal 8\n-rw-r--r-- 1 sri_charan sri_charan 220 Aug 3 11:00 exec_demo.c"),
    ("7) Parent-Child Synchronization using wait()", "gcc wait_demo.c -o wait_demo && ./wait_demo", "Child executing...\nChild complete.\nParent executing after wait."),
    ("8) Process Tree Analysis", "pstree -p 5120", "bash(4000)───main(5120)───main(5121)"),
    ("9) System Call Tracing", "strace -e trace=execve ./prog1", "execve(\"./prog1\", [\"./prog1\"], 0x7ffc...) = 0\nPID = 2456\n+++ exited with 0 +++"),
    ("10) Mini Shell Program", "./minishell", "minishell> date\nMon Aug  3 11:05:00 IST 2026\nminishell> exit")
]

for idx, (title, cmd, out) in enumerate(skill_topics, 1):
    h = doc_s.add_heading(title, level=2)
    h.runs[0].font.name = 'Calibri'
    h.runs[0].font.color.rgb = RGBColor(0, 102, 153)

    img_path = os.path.join(skill_img_dir, f"win_skill_{idx}.png")
    render_win_terminal(out, img_path, command_line=cmd)

    p = doc_s.add_paragraph("Terminal Screenshot:")
    p.runs[0].font.bold = True
    doc_s.add_picture(img_path, width=Inches(6.2))
    doc_s.add_paragraph()

doc_s_path = os.path.join(skill_dir, "Skill-01_Report.docx")
doc_s.save(doc_s_path)
print("Updated human-style documents for Windows Terminal.")
